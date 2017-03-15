# -*- coding: utf-8 -*-

from PyQt5.QtWidgets import QApplication, QTableWidgetItem, QMainWindow, QMessageBox, QTableWidget
from PyQt5 import QtCore
from settings import *
import json
from turn_assignment_ui import Ui_TurnAssignment


class TurnAssignment(QMainWindow, Ui_TurnAssignment):
    def __init__(self, parent=None):
        super(TurnAssignment, self).__init__(parent)
        QMainWindow.setWindowFlags(self, QtCore.Qt.WindowCloseButtonHint)  # 禁止最大化和最小化按钮

        self.setupUi(self)
        self.window_resize()

        self.btnConfig.clicked.connect(self.config)
        self.btnView.clicked.connect(self.view)
        self.btnGen.clicked.connect(self.gen)
        self.btnExport.clicked.connect(self.export)
        self.btnSave.clicked.connect(self.save_info)

        self.init_tabs_title()

        self.year = 2016      # 获取当前年份
        self.month = 11       # 获取当前月份
        self.cal = []         # 建立空的月历列表
        self.days = 30        # 初始化每月的天数为30

        self.cb_list = [self.cbMon, self.cbTus, self.cbWed, self.cbThu, self.cbFri, self.cbSat, self.cbSun]
        self.week_list = WEEKDAYS  # ["星期天","星期一", "星期二","星期三","星期四","星期五","星期六"]
        self.students = STUDENTS
        self.choice_date_list = []

        self.willing_dict = {}
        self.load_json_file()
        self.init_cbb_student()
        self.init_tab_config()
        self.info_table_gen()  # 生成初始信息表

        self.calender_gen()
        self.day_num_first_week = len(self.cal[0]) - self.cal[0].count("")
        self.duty_table = {}
        self.turn_table_gen()
        self.tabs.setCurrentIndex(0)  # 设置打开窗口后的显示的第一个tab页

        self.show()

    def init_tabs_title(self):
        self.tabs.setTabText(self.tabs.indexOf(self.tabConfig), "配置信息")
        self.tabs.setTabText(self.tabs.indexOf(self.tabShow), "信息表")
        self.tabs.setTabText(self.tabs.indexOf(self.tabTable), "值日表")

    def init_cbb_student(self):
        # 初始化学生选择comoBox
        self.cbbStudents.clear()
        for item in self.students:
            self.cbbStudents.addItem(item)

    def init_tab_config(self):

        self.cbMon.stateChanged['int'].connect(self.set_time)
        self.cbFri.stateChanged['int'].connect(self.set_time)
        self.cbSat.stateChanged['int'].connect(self.set_time)
        self.cbSun.stateChanged['int'].connect(self.set_time)
        self.cbThu.stateChanged['int'].connect(self.set_time)
        self.cbTus.stateChanged['int'].connect(self.set_time)
        self.cbWed.stateChanged['int'].connect(self.set_time)
        self.btnClean.clicked.connect(self.clear_choice)
        self.btnSubmit.clicked.connect(self.submit)
        self.calendar.clicked['QDate'].connect(self.set_time)
        self.calendar.activated['QDate'].connect(self.date_choice)
        self.cbbStudents.activated['int'].connect(self.student_selected)

        self.cbbMonth.setCurrentIndex(10)  # 初始化为11月
        self.cbbYear.setCurrentIndex(0)    # 初始化为2016年

        self.year = int(self.cbbYear.currentText())      # 获取当前年份
        self.month = int(self.cbbMonth.currentText())    # 获取当前月份

        self.calendar.setCurrentPage(self.year, self.month)       # 设置当前的年份和月份
        self.cbSat.setEnabled(False)
        self.cbSat.setVisible(False)     # 星期六不可见
        self.cbSun.setEnabled(False)
        self.cbSun.setVisible(False)     # 星期天不可见

        student = self.cbbStudents.currentText()    # 获取学生姓名
        willing = self.willing_dict[student]        # 获取学生家长的意愿

        self.choice_date_list = self.willing_dict[student][-1].split("'")   # 将第一个学生的日期信息写入日期选择列表

        for i, obj in enumerate(self.cb_list[:5]):
            obj.setChecked(willing[i] == "Y")

        self.cbbMonth.activated['int'].connect(self.update_calendar)
        self.cbbYear.activated['int'].connect(self.update_calendar)

    def update_calendar(self):
        self.year = int(self.cbbYear.currentText())
        self.month = int(self.cbbMonth.currentText())
        self.calendar.setCurrentPage(self.year, self.month)       # 设置当前的年份和月份
        self.calender_gen()
        self.day_num_first_week = len(self.cal[0]) - self.cal[0].count("")

    def load_json_file(self):
        try:
            f = open(WILLING_JSON, "r")
        except IOError:
            for student in self.students:
                self.willing_dict[student] = [""] * 6  # 初始化值日意愿表
            with open(WILLING_JSON, "w") as f:
                json.dump(self.willing_dict, f)
        else:
            self.willing_dict = json.load(f)
            f.close()

    def window_resize(self):
        self.resize(800, 600)
        self.tabs.setGeometry(QtCore.QRect(20, 30, 760, 550))
        self.horizontalLayout_5.setContentsMargins(80, 0, 0, 0)
        self.tableInfo.setGeometry(QtCore.QRect(0, 15, 760, 500))
        self.tableTurn.setGeometry(QtCore.QRect(0, 15, 760, 500))
        self.gridLayout.setContentsMargins(10, 10, 10, 10)

    def info_table_gen(self):
        self.tableInfo.setRowCount(len(self.students))  # 设置表格总行数
        set_info_item = self.tableInfo.setItem
        for index, item in enumerate(self.students):
            # self.tableInfo.setItem(index, 0, QTableWidgetItem(item))
            set_info_item(index, 0, QTableWidgetItem(item))
            for i in range(1, 7):
                # self.tableInfo.setItem(index, i, QTableWidgetItem(self.willing_dict[item][i - 1]))
                set_info_item(index, i, QTableWidgetItem(self.willing_dict[item][i - 1]))
        self.tableInfo.setEditTriggers(QTableWidget.NoEditTriggers)

    def turn_table_gen(self):
        self.tableTurn.setRowCount(3 * len(self.cal))
        set_turn_item = self.tableTurn.setItem
        for row_index, value_row in enumerate(self.cal):
            for col_index, value_cell in enumerate(value_row):
                # self.tableTurn.setItem(3 * row_index, col_index, QTableWidgetItem(value_cell))
                set_turn_item(3 * row_index, col_index, QTableWidgetItem(value_cell))

        # 隐藏纵向标题
        self.tableTurn.setVerticalHeaderLabels([""] * self.tableTurn.rowCount())
        self.tableTurn.setEditTriggers(QTableWidget.NoEditTriggers)

    def calender_gen(self):
        from calendar import month
        cal = month(self.year, self.month).split("\n")[2:-1]
        cal = [item.split() for item in cal]
        pack_size = 8 - len(cal[0])
        for i in range(pack_size):
            cal[0].insert(0, "")
        cal_1 = []
        for item in cal:
            cal_1 += item
        self.days = len(cal_1) - cal_1.count("")
        self.cal = []
        cal_append = self.cal.append
        for i in range((len(cal_1) // 7) + 1):
            cal_append(cal_1[7 * i: 7 * i + 7])

    def student_selected(self, index):
        student = STUDENTS[index]
        self.clear_label_in_config()
        self.unchecked_checkbox_in_config()
        self.load_checkbox_state_in_config(student)
        self.load_date_choice(student)

    def clear_label_in_config(self):
        self.lbWeekday.setText(MSG_WEEKDAY)
        self.lbDate.setText(MSG_DATE)

    def load_checkbox_state_in_config(self, student):
        for index, value in enumerate(self.willing_dict[student][:-1]):
            self.cb_list[index].setChecked(value == "Y")

    def unchecked_checkbox_in_config(self):
        for obj in self.cb_list:
            obj.setChecked(False)

    def load_date_choice(self, student):
        date_str = self.willing_dict[student][-1]
        self.lbDate.setText(MSG_DATE + date_str)
        self.choice_date_list = date_str.split(",")

    def set_label_date_text(self):
        self.lbDate.clear()
        msg_date = ",".join(self.choice_date_list)
        self.lbDate.setText(MSG_DATE + msg_date)

    def set_label_weekday_text(self):
        self.lbWeekday.clear()
        cb_list = self.cb_list
        msg_weekday = ",".join([self.week_list[1:][index]
                                for index, item in enumerate(cb_list) if item.checkState() == 2])
        self.lbWeekday.setText(MSG_WEEKDAY + msg_weekday)

    def set_time(self):
        self.set_label_date_text()
        self.set_label_weekday_text()

    def clear_choice(self):
        reply = QMessageBox.question(QMessageBox(), "警告！", "确定将所有输入清除？", QMessageBox.Yes, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.clear_label_in_config()
            self.choice_date_list = []
            self.unchecked_checkbox_in_config()

    def submit(self):
        reply = QMessageBox.question(QMessageBox(), "确认！", "信息将被添加到值日意愿信息表中。", QMessageBox.Yes, QMessageBox.No)
        if reply == QMessageBox.Yes:
            weekday_choice = [self.week_list[index+1] for index, item in enumerate(self.cb_list) if
                              item.checkState() != 0]
            student = self.cbbStudents.currentText()
            self.update_willing_dict(student, weekday_choice)
            self.write_date_to_info_table(student)
            self.write_weekday_to_info_table(student, weekday_choice)
        else:
            pass

    def update_willing_dict(self, student, weekday_choice):
        for index, day in enumerate(self.week_list[1:-1]):
            if day in weekday_choice:
                # 工作日增加到轮值字典中
                self.willing_dict[student][index] = "Y"
            else:
                self.willing_dict[student][index] = ""
        # 具体日期增加到轮值字典中
        self.willing_dict[student][-1] = (",".join(self.choice_date_list))

    def update_willing_json(self):
        with open(WILLING_JSON, "w") as f:
            json.dump(self.willing_dict, f)

    def write_weekday_to_info_table(self, student, weekday_choice):
        for index, day in enumerate(self.week_list[1:-1]):
            # 清空相应学生对应的行
            self.tableInfo.setItem(self.students.index(student),
                                   index + 1, QTableWidgetItem(""))
            if day in weekday_choice and student:
                # 工作日增加到表格相应单元中
                self.tableInfo.setItem(self.students.index(student),
                                       index + 1, QTableWidgetItem("Y"))

    def write_date_to_info_table(self, student):
        if student:
            # 具体日期增加到相应表格单元中
            self.tableInfo.setItem(self.students.index(student), 6, QTableWidgetItem(",".join(self.choice_date_list)))

    def date_choice(self):
        choice_date = self.calendar.selectedDate()
        choice_date_str = "-".join([str(item) for item in choice_date.getDate()[1:]])
        if choice_date_str not in self.choice_date_list:
            self.choice_date_list.append(choice_date_str)

    def config(self):
        self.tabs.setCurrentIndex(0)

    def view(self):
        self.tabs.setCurrentIndex(1)

    def save_info(self):
        self.update_willing_json()

    def gen(self):
        self.calender_gen()
        self.turn_table_gen()
        self.turn_assignment()
        self.tabs.setCurrentIndex(2)

    def turn_assignment(self):
        # 清空值日表
        for row in range(self.tableTurn.rowCount()):
            for column in range(self.tableTurn.columnCount()):
                self.tableTurn.setItem(row, column, None)
        self.turn_table_gen()
        # 安排值日顺序
        self.assignment_date()  # 优先按照具体日期进行安排
        self.assignment_weekday()  # 默认分配两次
        self.assignment_weekday()

    def assignment_date(self):
        non_assignment = self.assignment()
        self.assignment(first=False, source_dict=non_assignment)  # 二次分配

    def assignment(self, first=True, source_dict=None):
        non_assignment = {}
        if first:
            source_dict = self.willing_dict
        day_num_first_week = self.day_num_first_week

        if not source_dict:
            return

        for key, value in source_dict.items():
            if not value[-1]:
                continue

            if first:
                dates = [int(x[3:]) for x in value[-1].split(",") if x]
                dates.sort()
            else:
                dates = value

            for date in dates:
                row_index = 3 * ((date + 7 - day_num_first_week - 1) // 7) + 1
                column_index = (date + 7 - day_num_first_week - 1) % 7
                # 获取相应单元的内容，如果为空(None)则写入新内容（分配）
                item = self.tableTurn.item(row_index, column_index)
                if not item:
                    self.tableTurn.setItem(row_index, column_index, QTableWidgetItem(key))
                    if len(dates) > 1 and first:
                        non_assignment[key] = dates[1:]
                    break
                elif not (self.tableTurn.item(row_index + 1, column_index)):
                    self.tableTurn.setItem(row_index + 1, column_index, QTableWidgetItem(key))
                    if len(dates) > 1 and first:
                        non_assignment[key] = dates[dates.index(date) + 1:]
                    break
        return non_assignment

    def assignment_weekday(self):
        for key, value in self.willing_dict.items():
            weekdays_choice = value[:-1]
            if weekdays_choice.count("Y") == 0:
                continue
            weekdays_choice.insert(0, "")
            weekdays_choice.append("")

            assignment_list = []
            for _,item in enumerate(self.cal):
                assignment_item = []
                for i in range(len(item)):
                    assignment_item.append(key if (weekdays_choice[i] == "Y") and (item[i] != "") else "")
                assignment_list.append(assignment_item)

            for index, item in enumerate(assignment_list):
                assigned = False
                for index1, value1 in enumerate(item):
                    if not value1:
                        continue
                    if not self.tableTurn.item(3*index+1, index1):
                        self.tableTurn.setItem(3*index+1, index1, QTableWidgetItem(value1))
                        assigned = True
                        break
                    elif not self.tableTurn.item(3*index+2, index1):
                        self.tableTurn.setItem(3*index+2, index1, QTableWidgetItem(value1))
                        assigned = True
                        break
                if assigned:
                    break





            '''
            for index, weekday_choice in enumerate(weekdays_choice):
                if weekday_choice == "Y":
                    for i in range(len(self.cal)):
                        if (i == 0 and not self.cal[0][index+1]) or i > 0:                   # 第一行因为不总是从星期天开始需要妥善处理
                            item = self.tableTurn.item(3 * i, index + 1)
                            if not item:
                                self.tableTurn.setItem(3 * i, index + 1, QTableWidgetItem(key))
                                break
                            elif not (self.tableTurn.item(3 * i + 1, index + 1, )):
                                self.tableTurn.setItem(3 * i + 1, index, QTableWidgetItem(key))
                                break
        '''
        '''
        for key in self.willing_dict.keys():
            for index, willing in enumerate(self.willing_dict[key][:-1]):
                if willing == "Y":
                    for i in range(len(self.cal)):
                        item = self.tableTurn.item(3 * i, index + 1)
                        if not item:
                            self.tableTurn.setItem(3 * i, index + 1, QTableWidgetItem(key))
                            break
                        elif not (self.tableTurn.item(3 * i + 1, index + 1, )):
                            self.tableTurn.setItem(3 * i + 1, index, QTableWidgetItem(key))
                            break
        '''
    def export_to_csv(self, filename):
        """
        导出到csv文件
        return: None
        """
        with open(filename, "w") as f:
            for row in range(self.tableTurn.rowCount()):
                for column in range(self.tableTurn.columnCount()):
                    if self.tableTurn.item(row, column):
                        content = self.tableTurn.item(row, column).text()
                        f.write(content + ",")
                    else:
                        f.write(",")
                f.write("\n")

    def export_to_docx(self, filename):
        """
        导出成Word文件
        return: None
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT

        doc = Document('./templates/default.docx')
        section = doc.sections[-1]
        # 页面设为横向，下面四个语句缺一不可
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = new_height
        section.page_width = new_width

        from docx.shared import Pt
        style = doc.styles['Normal']
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(13)

        # p = doc.add_paragraph("")
        p = doc.add_paragraph("一年11班值日轮值表（{}年{}月）".format(self.year, self.month), )
        p.bold = True
        # p.add_run("一年11班值日轮值表（{}年{}月）".format(self.year, self.month), ).bold = True
        paragraph_format = p.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER     # 文字在页面居中

        table = doc.add_table(rows=self.tableTurn.rowCount() + 1,
                              cols=self.tableTurn.columnCount(),
                              style="Light Grid")

        table.alignment = WD_TABLE_ALIGNMENT.CENTER     # 表格在页面剧中

        header_cell = table.rows[0].cells
        for i in range(len(self.week_list)):
            header_cell[i].text = self.week_list[i]
        for row_index in range(1, self.tableTurn.rowCount() + 1):
            row_cell = table.rows[row_index].cells
            for col_index in range(self.tableTurn.columnCount()):
                if self.tableTurn.item(row_index, col_index):
                    row_cell[col_index].text = self.tableTurn.item(row_index, col_index).text()
                else:
                    row_cell[col_index].text = ""

        # 合并最后一行
        start_cell = table.cell(self.tableTurn.rowCount(), 0)
        end_cell = table.cell(self.tableTurn.rowCount(), self.tableTurn.columnCount()-1)
        start_cell.merge(end_cell)
        info = ("值日时间：19:10之后（除周三和周五外）、17:20之后（周三）、11:50之后（周五）\n"
                "值日内容：扫地、擦桌椅、擦门窗、擦储物柜、拖地、摆桌椅、整理图书"
                )
        table.cell(self.tableTurn.rowCount(), 0).text = info

        try:
            doc.save(filename)
        except IOError:
            QMessageBox.information(QMessageBox(), "警告！", "请将相关的Word文件关闭", QMessageBox.Ok)

    @staticmethod
    def snap_screen(filename):
        # 截屏并保存为jpg文件
        from PIL import ImageGrab
        img = ImageGrab.grab()
        img.save(filename)

    def export(self):
        """
        根据用户选择到处WORD文件、CSV文件，或保存为jpg文件
        return: None
        """
        from PyQt5.QtWidgets import QFileDialog

        filename, _ = QFileDialog.getSaveFileName(
            QFileDialog(),
            "文件保存", "./",
            "所有文件 (*.*);;csv文件 (*.csv);; Word文件 (*.docx);; jpg 文件 (*.jpg)")

        file_type = filename.split(".")[-1]

        if file_type == "docx":
            self.export_to_docx(filename)
        elif file_type == "csv":
            self.export_to_csv(filename)
        elif file_type == "jpg":
            self.snap_screen(filename)
        else:
            pass


if __name__ == '__main__':
    import sys

    app = QApplication(sys.argv)
    app.setApplicationName("值日表生成")

    turn = TurnAssignment()

    sys.exit(app.exec_())
